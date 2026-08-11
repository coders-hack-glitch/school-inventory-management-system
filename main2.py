from pathlib import Path
from tkinter import *
from tkinter import ttk, messagebox
import datetime
import os
from decimal import Decimal, InvalidOperation

from docx.shared import Inches, Pt, RGBColor

from db_config import get_connection
from ui_helpers import COLORS, card, styled_entry, button, configure_treeview_style

BASE_DIR = Path(__file__).resolve().parent


class Application:
    def __init__(self, master, *args, **kwargs):
        self.master = master
        self.master.title('New Sale')
        self.master.geometry('1120x700')
        self.master.configure(bg=COLORS['bg'])
        self.master.resizable(False, False)
        self.date = datetime.datetime.now().date()
        self.cart = []
        self.current_product = None
        configure_treeview_style()

        header = Frame(master, bg=COLORS['navy'], height=72); header.pack(fill='x'); header.pack_propagate(False)
        Label(header, text='NEW SALE', font=('Segoe UI', 22, 'bold'), bg=COLORS['navy'], fg='white').pack(side='left', padx=30, pady=16)
        Label(header, text=f"Today • {self.date.strftime('%d %b %Y')}", font=('Segoe UI', 10), bg=COLORS['navy'], fg='#CDE6F2').pack(side='right', padx=30)

        left = card(master, 515, 595); left.place(x=25, y=90)
        right = card(master, 555, 595); right.place(x=555, y=90)
        self.left = left; self.right = right

        Label(left, text='Find Product', font=('Segoe UI', 15, 'bold'), bg='white', fg=COLORS['text']).place(x=24, y=22)
        Label(left, text='Product ID', font=('Segoe UI', 10, 'bold'), bg='white', fg=COLORS['muted']).place(x=24, y=65)
        self.enteride = styled_entry(left); self.enteride.place(x=24, y=88, width=300, height=40)
        button(left, 'SEARCH', self.search_product, 'primary', 15).place(x=340, y=88, width=135, height=40)

        self.product_card = Frame(left, bg=COLORS['soft_blue'], highlightbackground='#CBE7F2', highlightthickness=1)
        self.product_card.place(x=24, y=148, width=451, height=110)
        self.productname = Label(self.product_card, text='No product selected', font=('Segoe UI', 16, 'bold'), bg=COLORS['soft_blue'], fg=COLORS['text'], anchor='w')
        self.productname.place(x=18, y=16, width=410)
        self.pprice = Label(self.product_card, text='Search by ID to view price and stock', font=('Segoe UI', 10), bg=COLORS['soft_blue'], fg=COLORS['muted'], anchor='w')
        self.pprice.place(x=18, y=50, width=410)
        self.stock_label = Label(self.product_card, text='', font=('Segoe UI', 9, 'bold'), bg=COLORS['soft_blue'], fg=COLORS['blue_dark'], anchor='w')
        self.stock_label.place(x=18, y=76, width=410)

        Label(left, text='Sale Details', font=('Segoe UI', 15, 'bold'), bg='white', fg=COLORS['text']).place(x=24, y=282)
        Label(left, text='Quantity', font=('Segoe UI', 10, 'bold'), bg='white', fg=COLORS['muted']).place(x=24, y=326)
        self.quantity_e = styled_entry(left); self.quantity_e.place(x=24, y=350, width=205, height=40)
        Label(left, text='Discount (₹)', font=('Segoe UI', 10, 'bold'), bg='white', fg=COLORS['muted']).place(x=250, y=326)
        self.discount_e = styled_entry(left); self.discount_e.place(x=250, y=350, width=225, height=40); self.discount_e.insert(END, '0')
        self.add_to_cart_btn = button(left, 'ADD TO CART', self.add_to_cart, 'success', 20)
        self.add_to_cart_btn.place(x=24, y=415, width=451, height=44)
        button(left, 'CLEAR CURRENT PRODUCT', self.clear_current, 'secondary', 20).place(x=24, y=470, width=451, height=38)
        Label(left, text='Tip: Search another product after adding to build a multi-item bill.', font=('Segoe UI', 9), bg='white', fg=COLORS['muted']).place(x=24, y=530)

        Label(right, text='Shopping Cart', font=('Segoe UI', 15, 'bold'), bg='white', fg=COLORS['text']).place(x=24, y=22)
        self.cart_count = Label(right, text='0 items', font=('Segoe UI', 9, 'bold'), bg='white', fg=COLORS['blue']); self.cart_count.place(x=455, y=28)
        table_frame = Frame(right, bg='white'); table_frame.place(x=20, y=62, width=515, height=300)
        cols = ('product', 'qty', 'discount', 'amount')
        self.tree = ttk.Treeview(table_frame, columns=cols, show='headings', style='Modern.Treeview', selectmode='browse')
        for col, heading, width in [('product','PRODUCT',215),('qty','QTY',60),('discount','DISC.',90),('amount','AMOUNT',120)]:
            self.tree.heading(col, text=heading); self.tree.column(col, width=width, anchor='center' if col != 'product' else 'w')
        scroll = ttk.Scrollbar(table_frame, orient='vertical', command=self.tree.yview); self.tree.configure(yscrollcommand=scroll.set)
        self.tree.pack(side='left', fill='both', expand=True); scroll.pack(side='right', fill='y')

        summary = Frame(right, bg='#F7FAFC', highlightbackground=COLORS['border'], highlightthickness=1); summary.place(x=20, y=382, width=515, height=100)
        Label(summary, text='Subtotal', font=('Segoe UI', 10), bg='#F7FAFC', fg=COLORS['muted']).place(x=18, y=14)
        Label(summary, text='Discount', font=('Segoe UI', 10), bg='#F7FAFC', fg=COLORS['muted']).place(x=18, y=42)
        self.subtotal_label = Label(summary, text='₹0.00', font=('Segoe UI', 10, 'bold'), bg='#F7FAFC', fg=COLORS['text']); self.subtotal_label.place(x=400, y=14)
        self.discount_label = Label(summary, text='₹0.00', font=('Segoe UI', 10, 'bold'), bg='#F7FAFC', fg=COLORS['danger']); self.discount_label.place(x=400, y=42)
        Label(summary, text='TOTAL', font=('Segoe UI', 13, 'bold'), bg='#F7FAFC', fg=COLORS['text']).place(x=250, y=68)
        self.total_label = Label(summary, text='₹0.00', font=('Segoe UI', 15, 'bold'), bg='#F7FAFC', fg=COLORS['success']); self.total_label.place(x=400, y=66)

        self.print_btn = button(right, 'PRINT BILL', self.generate_bill, 'primary', 20); self.print_btn.place(x=20, y=505, width=250, height=46); self.print_btn.config(state=DISABLED)
        button(right, 'CLEAR CART', self.clear_cart, 'danger', 20).place(x=285, y=505, width=250, height=46)
        self.enteride.focus()
        self.master.bind('<Return>', self.search_product)
        self.master.protocol('WM_DELETE_WINDOW', self.on_close)

    def search_product(self, *args):
        product_id = self.enteride.get().strip()
        if not product_id:
            messagebox.showerror('Search', 'Please enter a Product ID.', parent=self.master); return
        con = cursor = None
        try:
            con = get_connection(); cursor = con.cursor()
            cursor.execute('SELECT id,name,stock,cp,sp,vendor FROM inventory WHERE id=%s', (product_id,)); result = cursor.fetchone()
            if not result:
                self.current_product = None; self.productname.config(text='Product not found'); self.pprice.config(text=''); self.stock_label.config(text=''); return
            self.current_product = {'id':result[0], 'name':result[1], 'stock':result[2], 'cp':result[3], 'sp':result[4], 'vendor':result[5]}
            self.productname.config(text=str(result[1])); self.pprice.config(text=f"Selling price: ₹{Decimal(str(result[4])):.2f}")
            self.stock_label.config(text=f"Available stock: {result[2]}   •   Vendor: {result[5] or '—'}")
            self.quantity_e.delete(0, END); self.quantity_e.focus()
        except Exception as exc: messagebox.showerror('Database Error', str(exc), parent=self.master)
        finally:
            if cursor: cursor.close()
            if con: con.close()

    def add_to_cart(self):
        if not self.current_product:
            messagebox.showerror('Cart', 'Search for a product first.', parent=self.master); return
        try:
            qty = int(self.quantity_e.get().strip()); discount = Decimal(self.discount_e.get().strip() or '0')
            if qty <= 0 or discount < 0: raise ValueError
        except (ValueError, InvalidOperation):
            messagebox.showerror('Cart', 'Enter a valid positive quantity and non-negative discount.', parent=self.master); return
        already = sum(item['quantity'] for item in self.cart if item['id'] == self.current_product['id'])
        available = int(self.current_product['stock']) - already
        if qty > available:
            messagebox.showerror('Stock', f'Only {available} unit(s) are available.', parent=self.master); return
        unit = Decimal(str(self.current_product['sp'])); gross = unit * qty
        if discount > gross:
            messagebox.showerror('Discount', 'Discount cannot exceed the item amount.', parent=self.master); return
        self.cart.append({'id':self.current_product['id'], 'name':self.current_product['name'], 'quantity':qty, 'discount':discount, 'amount':gross-discount, 'unit_price':unit})
        self.refresh_cart(); self.clear_current()

    def clear_current(self):
        self.current_product = None; self.enteride.delete(0, END); self.quantity_e.delete(0, END); self.discount_e.delete(0, END); self.discount_e.insert(END, '0')
        self.productname.config(text='No product selected'); self.pprice.config(text='Search by ID to view price and stock'); self.stock_label.config(text=''); self.enteride.focus()

    def refresh_cart(self):
        for item in self.tree.get_children(): self.tree.delete(item)
        for index, item in enumerate(self.cart):
            self.tree.insert('', 'end', iid=str(index), values=(item['name'], item['quantity'], f"₹{item['discount']:.2f}", f"₹{item['amount']:.2f}"))
        subtotal = sum((i['unit_price']*i['quantity'] for i in self.cart), Decimal('0')); total = sum((i['amount'] for i in self.cart), Decimal('0')); discount = subtotal-total
        self.subtotal_label.config(text=f'₹{subtotal:.2f}'); self.discount_label.config(text=f'₹{discount:.2f}'); self.total_label.config(text=f'₹{total:.2f}')
        self.cart_count.config(text=f'{len(self.cart)} item(s)'); self.print_btn.config(state=NORMAL if self.cart else DISABLED)

    @staticmethod
    def _set_cell_shading(cell, fill):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.find(qn('w:shd'))
        if shd is None:
            shd = OxmlElement('w:shd')
            tcPr.append(shd)
        shd.set(qn('w:fill'), fill)

    @staticmethod
    def _set_cell_margins(cell, top=100, start=100, bottom=100, end=100):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = tcPr.first_child_found_in('w:tcMar')
        if tcMar is None:
            tcMar = OxmlElement('w:tcMar')
            tcPr.append(tcMar)
        for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
            node = tcMar.find(qn(f'w:{m}'))
            if node is None:
                node = OxmlElement(f'w:{m}')
                tcMar.append(node)
            node.set(qn('w:w'), str(v))
            node.set(qn('w:type'), 'dxa')

    @staticmethod
    def _set_table_borders(table, color='D8E2E8', size='6'):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = tblPr.first_child_found_in('w:tblBorders')
        if borders is None:
            borders = OxmlElement('w:tblBorders')
            tblPr.append(borders)
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            tag = f'w:{edge}'
            el = borders.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                borders.append(el)
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), size)
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color)

    @staticmethod
    def _set_cell_text(cell, text, *, bold=False, size=9, color='253746', align=None):
        from docx.shared import Pt, RGBColor
        cell.text = ''
        p = cell.paragraphs[0]
        if align is not None:
            p.alignment = align
        r = p.add_run(str(text))
        r.bold = bold
        r.font.name = 'Aptos'
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string(color)
        Application._set_cell_margins(cell, 90, 100, 90, 100)

    def generate_bill(self):
        if not self.cart:
            messagebox.showinfo('Bill', 'Cart is empty.', parent=self.master); return
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
            from docx.shared import Inches, Pt, RGBColor
        except ImportError:
            messagebox.showerror('Missing Package', 'Install python-docx with:\n\npython -m pip install python-docx', parent=self.master); return

        now = datetime.datetime.now()
        invoice_no = now.strftime('CCS-%Y%m%d-%H%M%S')
        invoice_dir = BASE_DIR / 'invoices' / str(self.date)
        invoice_dir.mkdir(parents=True, exist_ok=True)
        file_path = invoice_dir / f"Invoice_{now.strftime('%Y%m%d_%H%M%S')}.docx"

        subtotal = sum((i['unit_price'] * i['quantity'] for i in self.cart), Decimal('0'))
        total = sum((i['amount'] for i in self.cart), Decimal('0'))
        discount_total = subtotal - total

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

        # Document metadata
        doc.core_properties.title = f'Invoice {invoice_no}'
        doc.core_properties.subject = 'CCS School Atrauli Inventory Sale'
        doc.core_properties.author = 'CCS Inventory Management System'

        # Header
        header = doc.add_table(rows=1, cols=2)
        header.alignment = WD_TABLE_ALIGNMENT.CENTER
        header.autofit = False
        header.columns[0].width = Inches(4.7)
        header.columns[1].width = Inches(2.1)
        left_cell, right_cell = header.rows[0].cells
        self._set_cell_shading(left_cell, '123E5A')
        self._set_cell_shading(right_cell, '123E5A')
        self._set_cell_margins(left_cell, 170, 180, 170, 120)
        self._set_cell_margins(right_cell, 170, 120, 170, 180)
        p = left_cell.paragraphs[0]
        r = p.add_run('CCS INVENTORY')
        r.bold = True; r.font.name = 'Aptos Display'; r.font.size = Pt(24); r.font.color.rgb = RGBColor(255,255,255)
        p2 = left_cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(0)
        r = p2.add_run('INVENTORY MANAGEMENT SYSTEM')
        r.font.name = 'Aptos'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(205,230,242)
        p = right_cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run('SALES INVOICE')
        r.bold = True; r.font.name = 'Aptos'; r.font.size = Pt(13); r.font.color.rgb = RGBColor(255,255,255)
        p = right_cell.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(invoice_no); r.font.name = 'Aptos'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(205,230,242)
        self._set_table_borders(header, '123E5A', '0')

        # School/date information
        info = doc.add_table(rows=1, cols=2)
        info.alignment = WD_TABLE_ALIGNMENT.CENTER
        info.autofit = False
        info.columns[0].width = Inches(4.7); info.columns[1].width = Inches(2.1)
        c1, c2 = info.rows[0].cells
        self._set_cell_shading(c1, 'F3F7FA'); self._set_cell_shading(c2, 'F3F7FA')
        self._set_cell_text(c1, 'CCS School Atrauli (Aligarh)\nInventory & Sales Department', bold=True, size=9, color='253746')
        self._set_cell_text(c2, f'Date: {now.strftime("%d %b %Y")}\nTime: {now.strftime("%I:%M %p")}', bold=False, size=9, color='6D7B88', align=WD_ALIGN_PARAGRAPH.RIGHT)
        self._set_table_borders(info, 'D8E2E8', '4')

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

        # Items heading
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(5)
        r = p.add_run('ITEMS')
        r.bold = True; r.font.name = 'Aptos'; r.font.size = Pt(11); r.font.color.rgb = RGBColor(18,62,90)

        # Items table
        table = doc.add_table(rows=1, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        widths = [0.65, 2.55, 0.65, 1.05, 1.05, 1.15]
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)
        headers = ['ID', 'PRODUCT', 'QTY', 'UNIT PRICE', 'DISCOUNT', 'AMOUNT']
        for cell, h in zip(table.rows[0].cells, headers):
            self._set_cell_shading(cell, '123E5A')
            self._set_cell_text(cell, h, bold=True, size=8, color='FFFFFF', align=WD_ALIGN_PARAGRAPH.CENTER)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        for index, item in enumerate(self.cart):
            cells = table.add_row().cells
            values = [item['id'], item['name'], item['quantity'], f'Rs. {item["unit_price"]:.2f}', f'Rs. {item["discount"]:.2f}', f'Rs. {item["amount"]:.2f}']
            for col, (cell, val) in enumerate(zip(cells, values)):
                if index % 2 == 0:
                    self._set_cell_shading(cell, 'F7FAFC')
                align = WD_ALIGN_PARAGRAPH.LEFT if col == 1 else WD_ALIGN_PARAGRAPH.CENTER
                self._set_cell_text(cell, val, size=8.5, color='253746', align=align)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        self._set_table_borders(table, 'D8E2E8', '4')

        # Summary block
        spacer = doc.add_paragraph(); spacer.paragraph_format.space_after = Pt(1)
        summary = doc.add_table(rows=4, cols=2)
        summary.alignment = WD_TABLE_ALIGNMENT.RIGHT
        summary.autofit = False
        summary.columns[0].width = Inches(4.9); summary.columns[1].width = Inches(2.2)
        labels = [('Subtotal', f'Rs. {subtotal:.2f}'), ('Discount', f'- Rs. {discount_total:.2f}'), ('Net Amount', f'Rs. {total:.2f}'), ('Amount Payable', f'Rs. {total:.2f}')]
        for i, (label, value) in enumerate(labels):
            c1, c2 = summary.rows[i].cells
            fill = 'E8F7F0' if i == 3 else ('F3F7FA' if i % 2 == 0 else 'FFFFFF')
            self._set_cell_shading(c1, fill); self._set_cell_shading(c2, fill)
            self._set_cell_text(c1, label, bold=(i >= 2), size=9 if i < 2 else 10, color='253746', align=WD_ALIGN_PARAGRAPH.RIGHT)
            self._set_cell_text(c2, value, bold=True, size=9 if i < 2 else 11, color='1F9D67' if i == 3 else '253746', align=WD_ALIGN_PARAGRAPH.RIGHT)
        self._set_table_borders(summary, 'D8E2E8', '4')

        # Footer note
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(0)
        r = p.add_run('Thank you for using CCS Inventory Management System')
        r.bold = True; r.font.name = 'Aptos'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(18,62,90)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(0)
        r = p.add_run('Computer-generated invoice • No signature required')
        r.font.name = 'Aptos'; r.font.size = Pt(7.5); r.font.color.rgb = RGBColor(109,123,136)

        doc.save(file_path)

        try:
            os.startfile(str(file_path))
        except OSError as exc:
            messagebox.showwarning('Bill Created', f'Bill saved at:\n{file_path}\n\nWord could not be opened automatically.\n{exc}', parent=self.master)

        con = cursor = None
        try:
            con = get_connection(); cursor = con.cursor()
            for item in self.cart:
                cursor.execute('UPDATE inventory SET stock=stock-%s, totalcp=cp*stock, totalsp=sp*stock, assumed_profit=(sp*stock)-(cp*stock) WHERE id=%s AND stock >= %s', (item['quantity'], item['id'], item['quantity']))
                if cursor.rowcount == 0:
                    raise RuntimeError(f'Unable to update stock for product ID {item["id"]}.')
            con.commit()
        except Exception as exc:
            if con: con.rollback()
            messagebox.showwarning('Stock Update Warning', f'Bill was created, but stock was not updated.\n\n{exc}', parent=self.master)
        finally:
            if cursor: cursor.close()
            if con: con.close()

        messagebox.showinfo('Bill Created', f'Bill created successfully.\n\n{file_path}', parent=self.master)
        self.clear_cart()

    def clear_cart(self):
        self.cart.clear(); self.refresh_cart(); self.clear_current()
    def on_close(self): self.master.destroy()


if __name__ == '__main__':
    root = Tk(); Application(root); root.mainloop()
